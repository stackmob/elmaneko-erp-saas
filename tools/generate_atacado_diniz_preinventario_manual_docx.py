from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path("output/docx/ZY-ATACADODINIZ-2026-PREINVENTARIO-MAN-001.docx")
GREEN = "00B975"
DARK_GREEN = "00976A"
TEXT = "1A1A1A"
MUTED = "6B7280"
LABEL = "9CA3AF"
BORDER = "E5E7EB"
ALT = "F9FAFB"


def set_cell_shading(cell, fill):
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)
    shading.set(qn("w:val"), "clear")


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, **edges):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge, options in edges.items():
        tag = qn(f"w:{edge}")
        element = borders.find(tag)
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        for key, value in options.items():
            element.set(qn(f"w:{key}"), str(value))


def set_table_geometry(table, widths_inches):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    total_dxa = round(sum(widths_inches) * 1440)
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    indent = tbl_pr.first_child_found_in("w:tblInd")
    if indent is None:
        indent = OxmlElement("w:tblInd")
        tbl_pr.append(indent)
    indent.set(qn("w:w"), "120")
    indent.set(qn("w:type"), "dxa")
    for row in table.rows:
        for index, width in enumerate(widths_inches):
            row.cells[index].width = Inches(width)
            tc_w = row.cells[index]._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                row.cells[index]._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(round(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")


def set_run(run, size=None, bold=None, color=None, font="Arial"):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_bottom_border(paragraph, color=GREEN, size="12"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def add_page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, end])


def add_footer(section):
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    left = paragraph.add_run("Zaya IT - Gestão de Negócios | Manual Operacional | v1.0")
    set_run(left, 8, color=LABEL)
    paragraph.add_run(" " * 22)
    right = paragraph.add_run("Página ")
    set_run(right, 8, color=LABEL)
    add_page_field(paragraph)


def add_cover_footer(section):
    footer = section.first_page_footer
    table = footer.add_table(rows=1, cols=1, width=Inches(8.27))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_geometry(table, [6.5])
    cell = table.cell(0, 0)
    set_cell_shading(cell, DARK_GREEN)
    cell.text = ""
    cell.height = Inches(0.22)


def configure_document(document):
    section = document.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0.71)
    section.bottom_margin = Inches(0.87)
    section.left_margin = Inches(0.71)
    section.right_margin = Inches(0.71)
    section.header_distance = Inches(0.3)
    section.footer_distance = Inches(0.35)
    section.different_first_page_header_footer = True

    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(10)
    normal.font.color.rgb = RGBColor.from_string(TEXT)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name in ("List Bullet", "List Number"):
        style = document.styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.font.size = Pt(10)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25
    add_footer(section)
    add_cover_footer(section)


def add_text(paragraph, text, bold_phrases=()):
    cursor = 0
    for phrase in bold_phrases:
        position = text.find(phrase, cursor)
        if position == -1:
            continue
        if position > cursor:
            set_run(paragraph.add_run(text[cursor:position]))
        set_run(paragraph.add_run(phrase), bold=True)
        cursor = position + len(phrase)
    if cursor < len(text):
        set_run(paragraph.add_run(text[cursor:]))


def add_section(document, title):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(16)
    paragraph.paragraph_format.space_after = Pt(10)
    paragraph.paragraph_format.keep_with_next = True
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), GREEN)
    paragraph._p.get_or_add_pPr().append(shading)
    run = paragraph.add_run(title.upper())
    set_run(run, 11, bold=True, color="FFFFFF")
    return paragraph


def add_subsection(document, title):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(12)
    paragraph.paragraph_format.space_after = Pt(7)
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(title)
    set_run(run, 11.5, bold=True, color=GREEN)
    add_bottom_border(paragraph)
    return paragraph


def add_callout(document, label, text, kind="good"):
    settings = {
        "good": ("F0FDF4", GREEN, GREEN),
        "warning": ("FFFBEB", "F59E0B", "B45309"),
        "risk": ("FEF2F2", "EF4444", "B91C1C"),
    }
    fill, border, label_color = settings[kind]
    table = document.add_table(rows=1, cols=1)
    set_table_geometry(table, [6.5])
    cell = table.cell(0, 0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=120, start=180, bottom=100, end=180)
    set_cell_border(cell, left={"val": "single", "sz": "18", "color": border})
    label_paragraph = cell.paragraphs[0]
    label_paragraph.paragraph_format.space_after = Pt(4)
    label_run = label_paragraph.add_run(label.upper())
    set_run(label_run, 7.5, bold=True, color=label_color)
    content = cell.add_paragraph()
    content.paragraph_format.space_after = Pt(0)
    add_text(content, text, ("WMS > Inventário > Histórico de Ajuste de Estoque", "OK", "zero"))
    document.add_paragraph().paragraph_format.space_after = Pt(2)


def add_matrix(document, headers, rows, widths):
    table = document.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    header_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        cell = header_cells[index]
        set_cell_shading(cell, GREEN)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        run = cell.paragraphs[0].add_run(header)
        set_run(run, 8.5, bold=True, color="FFFFFF")
    for row_index, row in enumerate(rows):
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cell = cells[index]
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if row_index % 2 == 1:
                set_cell_shading(cell, ALT)
            set_cell_border(cell, top={"val": "single", "sz": "4", "color": BORDER}, bottom={"val": "single", "sz": "4", "color": BORDER}, left={"val": "single", "sz": "4", "color": BORDER}, right={"val": "single", "sz": "4", "color": BORDER})
            run = cell.paragraphs[0].add_run(value)
            set_run(run, 8.6, bold=(index == 0))
    document.add_paragraph().paragraph_format.space_after = Pt(2)


def add_numbered(document, items):
    for item in items:
        paragraph = document.add_paragraph(style="List Number")
        add_text(paragraph, item, ("Contagem de Estoque", "WMS > Inventário > Histórico de Ajuste de Estoque"))


def add_bullets(document, items):
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        add_text(paragraph, item, ("OK", "Histórico de Ajuste de Estoque"))


def build_document():
    document = Document()
    configure_document(document)

    for _ in range(6):
        document.add_paragraph()
    brand = document.add_paragraph()
    brand.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(brand.add_run("ZAYA IT"), 28, bold=True, color=GREEN)
    tag = document.add_paragraph()
    tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(tag.add_run("GESTÃO DE NEGÓCIOS"), 9, color=LABEL)
    rule = document.add_paragraph()
    rule.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(rule.add_run("━"), 18, color=GREEN)
    doctype = document.add_paragraph()
    doctype.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(doctype.add_run("MANUAL OPERACIONAL"), 8, color=GREEN)
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(title.add_run("Pré-Inventário WMS"), 22, bold=True, color=TEXT)
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(subtitle.add_run("Verificação do Cockpit, condução da contagem e acompanhamento dos ajustes"), 10, color=MUTED)
    badge = document.add_table(rows=1, cols=1)
    set_table_geometry(badge, [1.8])
    badge.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_cell_shading(badge.cell(0, 0), GREEN)
    badge.cell(0, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(badge.cell(0, 0).paragraphs[0].add_run("ATACADO DINIZ"), 7.5, bold=True, color="FFFFFF")
    document.add_paragraph()
    metadata = document.add_table(rows=5, cols=2)
    set_table_geometry(metadata, [1.7, 4.8])
    values = [("CLIENTE", "Atacado Diniz"), ("PROJETO", "Pré-Inventário WMS"), ("VERSÃO", "1.0"), ("EMISSÃO", date.today().strftime("%d/%m/%Y")), ("PÚBLICO", "Gestão logística, líderes de inventário e usuários-chave")]
    for row, (label, value) in zip(metadata.rows, values):
        set_cell_margins(row.cells[0]); set_cell_margins(row.cells[1])
        set_cell_border(row.cells[0], bottom={"val": "single", "sz": "4", "color": BORDER})
        set_cell_border(row.cells[1], bottom={"val": "single", "sz": "4", "color": BORDER})
        set_run(row.cells[0].paragraphs[0].add_run(label), 7.5, bold=True, color=GREEN)
        set_run(row.cells[1].paragraphs[0].add_run(value), 9.5, color=TEXT)
    document.add_page_break()

    add_section(document, "1. Objetivo e escopo")
    p = document.add_paragraph()
    add_text(p, "Este manual orienta a equipe do Atacado Diniz na preparação do inventário por meio do Cockpit Pré-Inventário WMS, na decisão de liberação da contagem e no acompanhamento dos ajustes realizados.")
    add_callout(document, "Premissa operacional", "O ambiente não utiliza coletor RF. Também não há tela padrão configurada para auditoria com recontagem nem para fechamento com emissão de notas fiscais. A recontagem é controlada operacionalmente e o acompanhamento posterior dos ajustes ocorre em WMS > Inventário > Histórico de Ajuste de Estoque.", "warning")
    add_callout(document, "Limite deste manual", "Não inicie a contagem enquanto houver indicador pendente no Cockpit. Quando o nome de uma tela for diferente no ambiente do usuário, interrompa a execução e solicite ao responsável o nome exato ou um print da tela antes de prosseguir.", "risk")

    add_section(document, "2. Verificação do Cockpit Pré-Inventário WMS")
    add_subsection(document, "2.1 Preparar a consulta")
    add_numbered(document, ["Abra o Cockpit Pré-Inventário WMS.", "Informe a empresa obrigatória e aplique os filtros necessários para a área que será inventariada.", "Quando aplicável, use os filtros de produto, rua, prédio inicial, prédio final, paridade e data inicial.", "Atualize a consulta antes de tomar qualquer decisão."])
    add_subsection(document, "2.2 Regra de liberação")
    add_callout(document, "Critério obrigatório", "O inventário só está liberado quando todos os sete indicadores do resumo estiverem em OK e com quantidade igual a zero. Qualquer valor diferente de zero indica processo em trânsito, pendência ou configuração que exige análise.")
    add_matrix(document, ["Indicador", "O que verificar", "Decisão"], [
        ("Tarefas pendentes/em execução", "Existência de tarefas ainda em andamento.", "Resolver antes da contagem."),
        ("Docas com saldo", "Saldo ainda existente em docas.", "Analisar e regularizar antes da contagem."),
        ("Docas configuradas como Ambos", "Docas que exigem análise de configuração.", "Validar com o responsável da operação."),
        ("Pedidos de venda pendentes de faturamento", "Pedidos que podem manter saldo lógico comprometido.", "Tratar antes da liberação."),
        ("Expedições pendentes", "Expedições não concluídas.", "Concluir ou tratar a pendência."),
        ("Recebimentos pendentes", "Recebimentos não concluídos.", "Concluir ou tratar a pendência."),
        ("Notas de compra não confirmadas", "Documentos de compra ainda não confirmados.", "Regularizar conforme processo interno."),
    ], [2.05, 2.45, 2.0])
    add_subsection(document, "2.3 Como tratar uma pendência")
    add_numbered(document, ["Identifique no resumo qual indicador está pendente.", "Abra o componente de detalhamento correspondente.", "Registre a pendência, o responsável e a decisão de tratamento conforme o processo interno.", "Depois da regularização, atualize o Cockpit e valide novamente o indicador.", "Libere a contagem somente após todos os indicadores retornarem a OK."])

    add_section(document, "3. Catálogo de telas usadas no processo")
    add_matrix(document, ["Tela ou componente", "Quando usar", "Resultado esperado"], [
        ("Cockpit Pré-Inventário - Resumo", "Antes de gerar ou iniciar a contagem.", "Todos os indicadores devem estar em situação OK e com quantidade zero."),
        ("Cockpit Pré-Inventário - Operação WMS", "Quando o resumo indicar pendência operacional.", "Identificar tarefas em aberto, docas com saldo e docas configuradas como Ambos."),
        ("Cockpit Pré-Inventário - Documentos Pendentes", "Quando o resumo indicar pendência comercial ou logística.", "Identificar pedidos de venda pendentes de faturamento, expedições, recebimentos e notas de compra não confirmadas."),
        ("Contagem de Estoque", "Durante a execução do inventário, conforme orientação do responsável pela operação.", "Registrar e acompanhar a contagem no fluxo habilitado para a empresa."),
        ("WMS > Inventário > Histórico de Ajuste de Estoque", "Após os ajustes autorizados e no acompanhamento do fechamento.", "Consultar o histórico dos ajustes realizados e manter a rastreabilidade da decisão operacional."),
    ], [2.1, 2.2, 2.2])
    add_callout(document, "Orientação ao usuário", "Este catálogo é intencionalmente funcional. Não utilize identificadores técnicos, nomes de tabelas ou nomes de campos para navegar no sistema. Se uma tela não corresponder ao nome acima, envie o nome exibido no menu ou um print da tela ao responsável pelo processo.", "warning")

    add_section(document, "4. Procedimento operacional")
    add_subsection(document, "4.1 Antes da contagem")
    add_numbered(document, ["Execute a verificação completa no Cockpit Pré-Inventário WMS.", "Trate todos os indicadores pendentes e confirme que o resumo está integralmente em OK.", "Registre a autorização do líder responsável para iniciar a contagem."])
    add_subsection(document, "4.2 Durante a contagem")
    add_numbered(document, ["Utilize a tela Contagem de Estoque de acordo com o fluxo habilitado para a empresa.", "Não presuma uso de coletor RF; siga o método de registro e conferência definido pela gestão logística.", "Ao identificar divergência, pause a decisão de ajuste e comunique o líder do inventário."])
    add_subsection(document, "4.3 Auditoria e recontagem sem rotina padrão")
    p = document.add_paragraph(); add_text(p, "A auditoria e a recontagem não são conduzidas por uma tela padrão neste ambiente. O líder do inventário deve registrar a solicitação, definir o responsável pela nova contagem e comparar o resultado com a primeira apuração antes de autorizar qualquer ajuste.")
    add_callout(document, "Controle recomendado", "Para cada recontagem, mantenha evidência da área conferida, data, responsável, motivo e decisão final. Não trate a divergência como concluída apenas porque a contagem foi repetida.", "warning")
    add_subsection(document, "4.4 Acompanhamento do ajuste")
    add_numbered(document, ["Após a autorização do ajuste, acesse WMS > Inventário > Histórico de Ajuste de Estoque.", "Consulte os registros relacionados ao inventário em análise.", "Confronte o resultado com a decisão aprovada pela liderança e mantenha a evidência no controle interno.", "Quando houver necessidade fiscal, siga o procedimento aprovado pelo setor fiscal. A emissão de notas fiscais não é executada por este manual."])
    add_section(document, "5. Checklist de encerramento")
    add_bullets(document, ["Todos os indicadores do Cockpit estão em OK e zerados antes da contagem.", "As pendências identificadas foram tratadas e registradas.", "A contagem foi autorizada pelo responsável da operação.", "As recontagens necessárias possuem evidência e decisão registrada.", "Os ajustes autorizados foram consultados no Histórico de Ajuste de Estoque.", "Demandas fiscais foram encaminhadas conforme o processo interno."])
    note = document.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text(note, "Documento elaborado a partir dos componentes do Cockpit Pré-Inventário WMS e das premissas operacionais confirmadas para o Atacado Diniz.")
    for run in note.runs:
        set_run(run, 8, color=LABEL)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    print(OUTPUT.resolve())


if __name__ == "__main__":
    build_document()
